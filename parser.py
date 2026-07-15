"""
Resume text extraction utilities.

Supports PDF (via pdfplumber) and DOCX (via python-docx) file formats.
"""

import os

import pdfplumber
from docx import Document


class ResumeParseError(Exception):
    """Raised when a resume file cannot be parsed."""


def extract_text_from_pdf(file_path: str) -> str:
    """Extract raw text from a PDF file using pdfplumber."""
    text_chunks = []
    try:
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text_chunks.append(page_text)
    except Exception as exc:
        raise ResumeParseError(f"Failed to read PDF file: {exc}") from exc

    text = "\n".join(text_chunks).strip()
    if not text:
        raise ResumeParseError("No readable text found in the PDF. It may be a scanned image.")
    return text


def extract_text_from_docx(file_path: str) -> str:
    """Extract raw text from a DOCX file using python-docx."""
    try:
        document = Document(file_path)
        paragraphs = [p.text for p in document.paragraphs if p.text.strip()]

        # Also pull text out of any tables in the document.
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text.strip())

        text = "\n".join(paragraphs).strip()
    except Exception as exc:
        raise ResumeParseError(f"Failed to read DOCX file: {exc}") from exc

    if not text:
        raise ResumeParseError("No readable text found in the DOCX file.")
    return text


def extract_resume_text(file_path: str) -> str:
    """Dispatch to the correct extractor based on file extension."""
    extension = os.path.splitext(file_path)[1].lower()

    if extension == ".pdf":
        return extract_text_from_pdf(file_path)
    elif extension == ".docx":
        return extract_text_from_docx(file_path)
    else:
        raise ResumeParseError(
            f"Unsupported file type '{extension}'. Please upload a PDF or DOCX file."
        )
